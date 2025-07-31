from typing import Optional
import asyncio
import tempfile
import shutil
import os
from urllib.parse import urlparse
import firebase_admin
from firebase_admin import credentials, auth, storage
from fastapi import FastAPI, HTTPException, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
import httpx
from docx import Document

from settings import settings
from helper_agents.transcriber_agent import transcriber_agent, TranscriberOutput
from helper_agents.summarizer_agent import summarizer_agent, SummarizerOutput
from agents import Runner


# Initialize FastAPI app
app = FastAPI(
    title="Audio Summarizer API",
    description="A REST API for transcribing and summarizing audio files using OpenAI Agents",
    version="1.0.0",
    contact={
        "name": "Jaime Santo",
        "email": "jaime.santo@gmail.com",
    },
    license_info={
        "name": "MIT",
    },
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=[settings.FRONTEND_URL],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Firebase Admin SDK
try:
    cred = credentials.Certificate(settings.FIREBASE_SERVICE_ACCOUNT_PATH)
    firebase_admin.initialize_app(cred, {
        'storageBucket': settings.FIREBASE_STORAGE_BUCKET
    })
    bucket = storage.bucket()
except Exception as e:
    print(f"Warning: Firebase initialization failed: {e}")
    bucket = None


# Pydantic models for request/response
class SummarizeRequest(BaseModel):
    audio_file_locator: str = Field(
        description="Firebase Storage URL for the audio file to transcribe (e.g., gs://bucket/audio/file.mp3)"
    )
    template_file_locator: str = Field(
        description="Firebase Storage URL for the template Word document (e.g., gs://bucket/templates/template.docx)"
    )


class SummarizeResponse(BaseModel):
    success: bool = Field(description="Whether the operation was successful")
    message: str = Field(description="Success or error message")
    report_file_locator: Optional[str] = Field(
        default=None, 
        description="Firebase Storage URL where the generated report is saved (only present on success)"
    )


async def validate_user_token(authorization: str = Header(...)) -> str:
    """Validate Firebase user token and return user_id"""
    if not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Invalid authorization header")
    
    token = authorization.replace("Bearer ", "")
    
    try:
        # Verify the Firebase token
        decoded_token = auth.verify_id_token(token)
        user_id = decoded_token['uid']
        return user_id
    except Exception as e:
        raise HTTPException(status_code=401, detail=f"Invalid token: {str(e)}")


def download_file_from_firebase(file_url: str, temp_path: str) -> str:
    """Download file from Firebase Storage to temporary location"""
    try:
        # Parse the Firebase Storage URL
        parsed_url = urlparse(file_url)
        if not parsed_url.scheme or not parsed_url.netloc:
            raise ValueError("Invalid Firebase Storage URL")
        
        # Extract the file path from the URL
        # Firebase Storage URLs are typically: gs://bucket-name/path/to/file
        if file_url.startswith('gs://'):
            # Handle gs:// URLs
            path_parts = file_url.replace('gs://', '').split('/', 1)
            if len(path_parts) != 2:
                raise ValueError("Invalid gs:// URL format")
            bucket_name, file_path = path_parts
        else:
            # Handle https:// URLs
            path_parts = parsed_url.path.split('/', 2)
            if len(path_parts) < 3:
                raise ValueError("Invalid Firebase Storage URL")
            bucket_name, file_path = path_parts[1], '/'.join(path_parts[2:])
        
        # Download the file
        blob = bucket.blob(file_path)
        blob.download_to_filename(temp_path)
        
        return temp_path
    except Exception as e:
        raise Exception(f"Failed to download file from Firebase: {str(e)}")


def upload_file_to_firebase(local_file_path: str, firebase_path: str) -> str:
    """Upload file to Firebase Storage"""
    try:
        blob = bucket.blob(firebase_path)
        blob.upload_from_filename(local_file_path)
        
        # Make the file publicly accessible (optional)
        blob.make_public()
        
        return f"gs://{settings.FIREBASE_STORAGE_BUCKET}/{firebase_path}"
    except Exception as e:
        raise Exception(f"Failed to upload file to Firebase: {str(e)}")


async def transcribe_audio(audio_file_path: str) -> str:
    """
    Transcribe audio file using OpenAI Agents SDK
    """
    try:
        # Call the transcriber agent
        result = await Runner.run( 
            transcriber_agent, 
            f"Transcribe the audio file at {audio_file_path}"
        )
        
        if isinstance(result.final_output, TranscriberOutput):
            assert len(result.final_output.transcript) > 0, "Transcript should not be empty"
            print(f"✓ Successfully transcribed via AI agent: {len(result.final_output.transcript)} characters")
        else:
            # Fallback for string output
            assert len(str(result.final_output)) > 0, "Transcript should not be empty"
            print(f"✓ Successfully transcribed via AI agent: {len(str(result.final_output))} characters")

        return result.final_output.transcript
        
    except Exception as e:
        raise Exception(f"Failed to transcribe audio: {str(e)}")


def extract_template_content(template_path: str) -> str:
    """Extract content from template Word document"""
    try:
        doc = Document(template_path)
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        print(f"✓ Successfully extracted template content: {len(content)} paragraphs")
        return "\n".join(content)

    except Exception as e:
        raise Exception(f"Failed to extract template content: {str(e)}")


async def summarize_transcript(transcript: str, template_content: str) -> str:
    """
    Generate report using OpenAI Agents SDK
    """
    try:
        result = await Runner.run(
            summarizer_agent,
            f"Generate a report based on the following transcript and template: {transcript}\n\nTemplate: {template_content}"
        )
        
        if isinstance(result.final_output, SummarizerOutput):
            assert len(result.final_output.summary) > 0, "Summary should not be empty"
            print(f"✓ Successfully generated summary via AI agent: {len(result.final_output.summary)} characters")
            return result.final_output.summary
        else:
            # Fallback for string output
            assert len(str(result.final_output)) > 0, "Summary should not be empty"
            print(f"✓ Successfully generated summary via AI agent: {len(str(result.final_output))} characters")
            return str(result.final_output)

    except Exception as e:
        raise Exception(f"Failed to summarize transcript: {str(e)}")


def save_report(summary: str, report_file_path: str) -> str:
    """Create Word document with summary"""
    try:
        doc = Document()
        doc.add_heading('Audio Transcription Report', 0)
                        
        # Add generated report
        doc.add_heading('Analysis Report', level=1)
        doc.add_paragraph(summary)
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_doc:
            doc.save(temp_doc.name)
            temp_file_path = temp_doc.name
        
        # Close the document to release the file handle
        del doc
        
        # Now move the file after the document is closed
        shutil.move(temp_file_path, report_file_path)
        print(f"✓ Successfully saved report: {report_file_path}")
        return report_file_path

    except Exception as e:
        raise Exception(f"Failed to save report: {str(e)}")


async def summarize_audio(
    audio_file_locator: str,
    template_file_locator: str,
    user_id: str
) -> str:
    """
    Transcribe audio file and generate report using OpenAI Agents SDK
    """
    temp_audio_file = None
    temp_template_file = None
    temp_report_file = None
    
    try:
        # Create temporary files
        temp_audio_file = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False)
        temp_template_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        temp_report_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        
        # Close the temporary files so they can be used by other functions
        temp_audio_file.close()
        temp_template_file.close()
        temp_report_file.close()
        
        print(f"\nDownloading audio file => ... ")
        download_file_from_firebase(audio_file_locator, temp_audio_file.name)
        
        print(f"\nDownloading template file => ... ")
        download_file_from_firebase(template_file_locator, temp_template_file.name)
        
        print(f"\nTranscribing audio => ... ")
        transcript = await transcribe_audio(temp_audio_file.name)
        
        print(f"\nExtracting template content => ... ")
        template_content = extract_template_content(temp_template_file.name)
        
        print(f"\nSummarizing transcript => ... ")
        summary = await summarize_transcript(transcript, template_content)
        
        print(f"\nSaving report => ... ")
        save_report(summary, temp_report_file.name)
        
        # Extract audio file name from the locator for the report name
        audio_file_name = os.path.basename(audio_file_locator.split('/')[-1])
        if not audio_file_name.endswith('.mp3'):
            audio_file_name += '.mp3'
        
        # Upload report to Firebase
        firebase_report_path = f"reports/{user_id}/{audio_file_name.replace('.mp3', '.docx')}"
        report_file_locator = upload_file_to_firebase(temp_report_file.name, firebase_report_path)
        
        print(f"\nUploading report to Firebase => ... ")
        print(f"✓ Successfully uploaded report: {report_file_locator}")
        
        return report_file_locator
        
    except Exception as e:
        raise Exception(f"Failed to summarize audio: {str(e)}")
    finally:
        # Clean up temporary files
        for temp_file in [temp_audio_file, temp_template_file, temp_report_file]:
            if temp_file and os.path.exists(temp_file.name):
                try:
                    os.unlink(temp_file.name)
                except:
                    pass


@app.post("/api/summarize", response_model=SummarizeResponse)
async def summarize_audio_endpoint(
    request: SummarizeRequest,
    user_id: str = Depends(validate_user_token)
):
    """
    Summarize audio file endpoint
    
    This endpoint transcribes an audio file and generates a summary report based on a template.
    
    **Process:**
    1. Downloads audio and template files from Firebase Storage
    2. Transcribes the audio using OpenAI Agents
    3. Extracts content from the template document
    4. Generates a summary report using AI
    5. Uploads the report back to Firebase Storage
    
    **Authentication:** Requires a valid Firebase user token in the Authorization header
    
    **Returns:** A JSON response with success status, message, and report file location
    """
    try:
        # Validate settings
        if not settings.validate():
            return SummarizeResponse(
                success=False,
                message="Server configuration error: Missing required environment variables"
            )
        
        # Process the audio summarization
        report_file_locator = await summarize_audio(
            request.audio_file_locator,
            request.template_file_locator,
            user_id
        )
        
        return SummarizeResponse(
            success=True,
            message="Audio summarization completed successfully",
            report_file_locator=report_file_locator
        )
        
    except Exception as e:
        return SummarizeResponse(
            success=False,
            message=f"Failed to summarize audio: {str(e)}"
        )


@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "message": "Audio Summarizer API is running"}


if __name__ == "__main__":
    import uvicorn
    
    # Validate settings before starting
    if not settings.validate():
        print("Error: Missing required environment variables")
        exit(1)
    
    uvicorn.run(
        "api:app",
        host=settings.API_HOST,
        port=settings.API_PORT,
        reload=True
    ) 