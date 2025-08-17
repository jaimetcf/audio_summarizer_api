import os
from typing import Optional, List
import asyncio
import argparse
from agents import Runner
from docx import Document
import tempfile
import shutil
from pydub import AudioSegment
import math

from settings import settings
from helper_agents.transcriber_agent_2 import transcriber_agent_2, Transcriber2Output
from helper_agents.summarizer_agent import summarizer_agent, SummarizerOutput


# Default values
default_audio_file_name = 'transformer-paper-introduction.mp3'
default_template_file_name = 'default_report_template.docx'

audio_files_folder = './data/audio_files/'
transcript_files_folder = './data/transcripts/'
template_files_folder = './data/report_templates/'
reports_folder = './data/reports/'



def parse_arguments():
    """Parse command line arguments"""
    parser = argparse.ArgumentParser(description='Audio Summarizer - Transcribe and summarize audio files')
    parser.add_argument(
        '--audio-file', 
        '-a',
        type=str, 
        default=default_audio_file_name,
        help=f'Audio file name (default: {default_audio_file_name})'
    )
    parser.add_argument(
        '--template-file', 
        '-t',
        type=str, 
        default=default_template_file_name,
        help=f'Template file name (default: {default_template_file_name})'
    )
    return parser.parse_args()


async def summarize_audio(
    audio_file_path: str,
    template_file_path: str,
):
    """
    Transcribe audio file and generate report using OpenAI Agents SDK
    """
    try:
        print(f"\nBreaking down audio file => ... ")
        chunk_files = break_down_audio_file(audio_file_path)
        print(f"✓ Successfully broken down audio file into {len(chunk_files)} chunks")
        print(f"  Chunk files: {chunk_files}")
        print('\n')

        
        print(f"\nTranscribing audio => ... ")
        all_transcripts = []
        
        for i, chunk_file in enumerate(chunk_files):
            print(f"\n\n  Transcribing chunk {i+1}/{len(chunk_files)}: {os.path.basename(chunk_file)}")
            transcript = await transcribe_audio(chunk_file)
            #print(f"\nTranscript: {i+1}/{len(chunk_files)}: {transcript}")
            all_transcripts.append(transcript)
        
        # Combine all transcripts
        full_transcript = "\n\n".join(all_transcripts)
        
        print(f"\nSaving transcript => ... ")
        audio_file_name = os.path.basename(audio_file_path)
        transcript_file_name = audio_file_name.replace('.mp3', '.txt')
        transcript_file_path = os.path.join(transcript_files_folder, transcript_file_name)        
        os.makedirs(transcript_files_folder, exist_ok=True)        
        with open(transcript_file_path, 'w', encoding='utf-8') as f:
            f.write(full_transcript)
        print(f"✓ Successfully saved transcript: {transcript_file_path}")
        
        print(f"\nExtracting template content => ... ")
        template_content = extract_template_content(template_file_path)
        
        print(f"\nSummarizing transcript => ... ")
        summary = await summarize_transcript(full_transcript, template_content)
        
        print(f"\nSaving report => ... ")
        save_report(summary, report_file_path)
        print('\n\n')
        
    except Exception as e:
        raise Exception(f"Failed to summarize audio: {str(e)}")

def break_down_audio_file(audio_file_path: str) -> List[str]:
    """
    Break down audio file into smaller chunks of 25MB or less
    
    Args:
        audio_file_path: Path to the original audio file
        
    Returns:
        List of paths to the smaller audio files
    """
    try:
        # Load the audio file
        audio = AudioSegment.from_mp3(audio_file_path)
        
        # Get file size in bytes
        file_size = os.path.getsize(audio_file_path)
        file_size_mb = file_size / (1024 * 1024)  # Convert to MB
        
        print(f"Original file size: {file_size_mb:.2f} MB")
        
        # If file is already 25MB or less, return the original file path
        if file_size_mb <= 25:
            print("✓ File is already within size limit (25MB or less)")
            return [audio_file_path]
        
        # Calculate number of chunks needed
        # n = file_size_mb / 25 + 1 (if there's a remainder)
        n_chunks = math.ceil(file_size_mb / 25)
        print(f"Breaking down into {n_chunks} chunks")
        
        # Calculate duration per chunk
        total_duration_ms = len(audio)
        chunk_duration_ms = total_duration_ms // n_chunks
        
        # Get base filename without extension
        base_name = os.path.splitext(os.path.basename(audio_file_path))[0]
        base_path = os.path.dirname(audio_file_path)
        
        chunk_files = []
        
        for i in range(n_chunks):
            # Calculate start and end times for this chunk
            start_time = i * chunk_duration_ms
            end_time = start_time + chunk_duration_ms if i < n_chunks - 1 else total_duration_ms
            
            # Extract the chunk
            chunk = audio[start_time:end_time]
            
            # Create chunk filename
            chunk_filename = f"{base_name}_part{i+1:02d}.mp3"
            #chunk_path = os.path.join(base_path, chunk_filename)
            chunk_path = base_path + '/' + chunk_filename
            
            # Export the chunk
            chunk.export(chunk_path, format="mp3")
            
            # Verify the chunk size
            chunk_size_mb = os.path.getsize(chunk_path) / (1024 * 1024)
            print(f"✓ Created chunk {i+1}/{n_chunks}: {chunk_filename} ({chunk_size_mb:.2f} MB)")
            
            chunk_files.append(chunk_path)
        
        return chunk_files
        
    except Exception as e:
        raise Exception(f"Failed to break down audio file: {str(e)}")

async def transcribe_audio(audio_file_path: str) -> str:
    """
    Transcribe audio file using OpenAI Agents SDK
    """
    try:
        # Call the transcriber agent
        result = await Runner.run( 
            transcriber_agent_2, 
            f"Transcribe the audio file at {audio_file_path}"
        )
        
        if isinstance(result.final_output, Transcriber2Output):
            # Convert the list of TranscriberParagraph objects to a single transcript string
            transcript_parts = []
            for paragraph in result.final_output.transcript:
                transcript_parts.append(f"{paragraph.speaker}: {paragraph.paragraph}")
            
            full_transcript = "\n\n".join(transcript_parts)
            assert len(full_transcript) > 0, "Transcript should not be empty"
            print(f"✓ Successfully transcribed via AI agent: {len(full_transcript)} characters")
            #print(f"  Sample transcript: {full_transcript[:100]}...")
            return full_transcript
        else:
            # Fallback for string output
            assert len(str(result.final_output)) > 0, "Transcript should not be empty"
            print(f"✓ Successfully transcribed via AI agent: {len(str(result.final_output))} characters")
            #print(f"  Sample transcript: {str(result.final_output)[:100]}...")
            return str(result.final_output)
        
    except Exception as e:
        raise Exception(f"Failed to transcribe audio: {str(e)}")


def extract_template_content(template_path: str) -> str:
    """Extract content from template Word document"""
    try:
        doc = Document(template_path)
        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        print(f"✓ Successfully extracted template content: {len(content)} paragraphs")
        #print(f"  Sample template content: {content[:100]}...")
        return "\n".join(content)

    except Exception as e:
        raise Exception(f"Failed to extract template content: {str(e)}")


async def summarize_transcript(transcript: str, template_content: str) -> str:
    """
    Generate report using OpenAI Agents SDK
    """
    try:
        result = await Runner.run(
            summarizer_agent,
            f"Generate a report based on the following transcript and template: {transcript}\n\nTemplate: {template_content}"
        )
        
        if isinstance(result.final_output, SummarizerOutput):
            assert len(result.final_output.summary) > 0, "Summary should not be empty"
            print(f"✓ Successfully generated summary via AI agent: {len(result.final_output.summary)} characters")
            #print(f"  Sample summary: {result.final_output.summary[:100]}...")
            return result.final_output.summary
        else:
            # Fallback for string output
            assert len(str(result.final_output)) > 0, "Summary should not be empty"
            print(f"✓ Successfully generated summary via AI agent: {len(str(result.final_output))} characters")
            #print(f"  Sample summary: {str(result.final_output)[:100]}...")
            return str(result.final_output)

    except Exception as e:
        raise Exception(f"Failed to summarize transcript: {str(e)}")


def save_report(summary: str, report_file_path: str) -> str:
    """Create Word document with summary"""
    try:
        doc = Document()
        doc.add_heading('Audio Transcription Report', 0)
                        
        # Add generated report
        doc.add_heading('Analysis Report', level=1)
        doc.add_paragraph(summary)
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_doc:
            doc.save(temp_doc.name)
            temp_file_path = temp_doc.name
        
        # Close the document to release the file handle
        del doc
        
        # Now move the file after the document is closed
        shutil.move(temp_file_path, report_file_path)
        print(f"✓ Successfully saved report: {report_file_path}")
        return report_file_path

    except Exception as e:
        raise Exception(f"Failed to save report: {str(e)}")


if __name__ == "__main__":
    # Parse command line arguments
    args = parse_arguments()
    
    # Set file paths based on CLI arguments
    audio_file_name = args.audio_file
    template_file_name = args.template_file
    
    audio_file_path = audio_files_folder + audio_file_name
    template_file_path = template_files_folder + template_file_name
    report_file_path = reports_folder + audio_file_name.replace('.mp3', '.docx')

    print('\nFiles that will be processed:    ')
    print(audio_file_path)
    print(template_file_path)

    print('\nReport file that will be saved:')
    print(report_file_path)

    print('\n--------------------------------------------------------------------\n')
    
    # Run all tests
    success = asyncio.run(summarize_audio(audio_file_path, template_file_path))    
    # Exit with appropriate code
    exit(0 if success else 1) 
